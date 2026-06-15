from docx import Document

doc = Document()
doc.add_heading('John Doe', 0)
doc.add_paragraph('john.doe@example.com | (123) 456-7890 | San Francisco, CA')

doc.add_heading('Education', level=1)
doc.add_paragraph('Bachelor of Science in Computer Science - Stanford University')

doc.add_heading('Experience', level=1)
doc.add_paragraph('Senior Software Engineer - Tech Corp (2020 - Present)')
doc.add_paragraph('Designed and built high-performance microservices using Python, FastAPI, and PostgreSQL. Integrated Docker container workflows and deployed services to AWS.')

doc.add_paragraph('Software Engineer - Startup Inc (2018 - 2020)')
doc.add_paragraph('Developed web applications using React, TypeScript, and Python. Configured databases using PostgreSQL and MongoDB.')

doc.add_heading('Projects', level=1)
doc.add_paragraph('CareerPilot AI: An automated resume intelligence parsing application.')

doc.add_heading('Certifications', level=1)
doc.add_paragraph('AWS Certified Solutions Architect - Associate')

doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, FastAPI, React, PostgreSQL, Docker, AWS, TypeScript, spaCy')

doc.save('/media/mohammad-ali/Ali/Clients/CareerPilot AI/test_resume.docx')
print("Successfully generated test_resume.docx!")
