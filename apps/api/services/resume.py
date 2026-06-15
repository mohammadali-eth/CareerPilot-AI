import json
import logging
import io
from datetime import datetime
from uuid import UUID
from typing import Dict, Any, List, Tuple

import httpx
from pypdf import PdfReader
from docx import Document
import spacy
from sqlalchemy.ext.asyncio import AsyncSession

from core.config import settings
from models.resume import Resume, ResumeScore, ATSReport
from repositories.resume import ResumeRepository, ResumeScoreRepository, ATSReportRepository

logger = logging.getLogger(__name__)

# Try to load spaCy model, download if missing
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logger.warning("spaCy 'en_core_web_sm' model not found. Downloading...")
    try:
        from spacy.cli import download
        download("en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")
    except Exception as e:
        logger.error(f"Failed to download spaCy model: {e}. Falling back to blank model.")
        nlp = spacy.blank("en")


class ResumeAnalyzerService:
    """
    Orchestrates file reading, spaCy classification, OpenAI structured analysis, and DB saves.
    """

    def __init__(self, db: AsyncSession):
        self.db = db
        self.resume_repo = ResumeRepository(db)
        self.score_repo = ResumeScoreRepository(db)
        self.ats_repo = ATSReportRepository(db)

    async def parse_pdf(self, file_bytes: bytes) -> str:
        """
        Extract text content from raw PDF file bytes using PyPDF.
        """
        text = ""
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error parsing PDF file: {e}")
            raise ValueError(f"Failed to parse PDF document: {str(e)}")
        return text.strip()

    async def parse_docx(self, file_bytes: bytes) -> str:
        """
        Extract text content from raw DOCX file bytes using python-docx.
        """
        text = ""
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            
            # Read paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            
            # Read table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {e}")
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")
        return text.strip()

    def _extract_local_fallback(self, raw_text: str) -> Dict[str, Any]:
        """
        Local heuristic analyzer using spaCy. Used when OpenAI API keys are absent.
        """
        doc = nlp(raw_text.lower())
        
        # 1. Simple Keyword dictionary matching for Skills
        common_skills = {
            "python", "javascript", "typescript", "react", "vue", "angular", "node", "express",
            "fastapi", "django", "flask", "postgresql", "mysql", "mongodb", "redis", "docker",
            "kubernetes", "aws", "gcp", "azure", "graphql", "rest", "git", "ci/cd", "terraform",
            "html", "css", "sass", "tailwind", "next.js", "nest.js", "go", "java", "c++", "rust"
        }
        
        found_skills = []
        for token in doc:
            if token.text in common_skills and token.text not in found_skills:
                found_skills.append(token.text)
                
        # 2. Basic Section matches
        sections = {"education": [], "experience": [], "certifications": [], "projects": []}
        lines = raw_text.split("\n")
        
        current_section = None
        for line in lines:
            cleaned = line.strip().lower()
            if not cleaned:
                continue
            
            # Identify section boundaries
            if "education" in cleaned or "university" in cleaned or "college" in cleaned:
                current_section = "education"
                continue
            elif "experience" in cleaned or "employment" in cleaned or "work history" in cleaned:
                current_section = "experience"
                continue
            elif "certification" in cleaned or "credentials" in cleaned:
                current_section = "certifications"
                continue
            elif "project" in cleaned or "portfolio" in cleaned:
                current_section = "projects"
                continue
                
            if current_section and len(sections[current_section]) < 10:
                sections[current_section].append(line.strip())

        # 3. Local Score Calculations
        has_edu = len(sections["education"]) > 0
        has_exp = len(sections["experience"]) > 0
        has_proj = len(sections["projects"]) > 0
        has_skills = len(found_skills) > 0
        
        structure_score = 40 + (15 if has_edu else 0) + (15 if has_exp else 0) + (15 if has_proj else 0) + (15 if has_skills else 0)
        content_score = min(40 + len(found_skills) * 3, 100)
        overall_score = int((structure_score + content_score) / 2)
        
        missing = ["kubernetes", "graphql", "typescript", "aws", "docker"]
        missing_keywords = [m for m in missing if m not in found_skills]
        
        ats_score = int(content_score - len(missing_keywords) * 2)
        ats_score = max(min(ats_score, 100), 10)

        suggestions = {
            "structure": [],
            "content": []
        }
        
        if not has_edu:
            suggestions["structure"].append("Add an education section to document academic history.")
        if not has_exp:
            suggestions["structure"].append("Add an experience section detailing your previous jobs.")
        if len(found_skills) < 5:
            suggestions["content"].append("Increase technical skills keyword density to satisfy search indexes.")

        return {
            "skills": [s.capitalize() for s in found_skills],
            "education": sections["education"],
            "experience": sections["experience"],
            "certifications": sections["certifications"],
            "projects": sections["projects"],
            "ats_score": ats_score,
            "overall_score": overall_score,
            "structure_score": structure_score,
            "content_score": content_score,
            "missing_keywords": missing_keywords,
            "formatting_issues": ["No major layout issues found"] if len(raw_text) > 300 else ["Short CV content"],
            "relevance_score": ats_score,
            "suggestions": suggestions
        }

    async def _call_openai_analysis(self, raw_text: str) -> Dict[str, Any]:
        """
        Call OpenAI to request structured resume metrics.
        """
        if not settings.OPENAI_API_KEY:
            logger.info("No OpenAI API key configured. Utilizing local spaCy fallback system.")
            return self._extract_local_fallback(raw_text)

        prompt = f"""
        You are a principal recruiting engineer and Applicant Tracking System (ATS) auditor.
        Parse and evaluate the following resume text. Return a strict JSON response.
        Do not output any introductory or concluding text, only raw valid JSON.

        JSON schema to match:
        {{
          "skills": ["skill1", "skill2"],
          "education": ["edu1", "edu2"],
          "experience": ["job1", "job2"],
          "certifications": ["cert1"],
          "projects": ["project1"],
          "ats_score": 85,
          "overall_score": 80,
          "structure_score": 90,
          "content_score": 70,
          "missing_keywords": ["GraphQL", "Kubernetes"],
          "formatting_issues": ["No contact information provided"],
          "relevance_score": 80,
          "suggestions": {{
            "structure": ["Add a summary"],
            "content": ["Include percentages in metrics"]
          }}
        }}

        Resume text:
        {raw_text}
        """

        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.post(
                    "https://api.openai.com/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {settings.OPENAI_API_KEY}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": settings.OPENAI_DEFAULT_MODEL,
                        "messages": [
                            {"role": "system", "content": "You are a professional recruiting analyzer. Output only JSON."},
                            {"role": "user", "content": prompt}
                        ],
                        "temperature": 0.1,
                        "response_format": {"type": "json_object"}
                    }
                )
                
                if response.status_code == 200:
                    result_json = response.json()
                    content = result_json["choices"][0]["message"]["content"]
                    return json.loads(content)
                else:
                    logger.error(f"OpenAI API error: {response.text}. Falling back to spaCy.")
                    return self._extract_local_fallback(raw_text)
        except Exception as e:
            logger.error(f"Failed to query OpenAI completions: {e}. Falling back to spaCy.")
            return self._extract_local_fallback(raw_text)

    async def analyze_and_store(
        self, user_id: UUID, filename: str, file_bytes: bytes, file_type: str
    ) -> Tuple[Resume, ResumeScore, ATSReport]:
        """
        Main runner: Parse file, analyze via OpenAI/spaCy, persist reports in DB.
        """
        # 1. Parse raw text
        if file_type == "application/pdf":
            raw_text = await self.parse_pdf(file_bytes)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            raw_text = await self.parse_docx(file_bytes)
        else:
            raise ValueError("Unsupported document format. Only PDF and DOCX documents are accepted.")

        if not raw_text:
            raise ValueError("Document appears to be empty or unreadable.")

        # 2. Run NLP & AI Extraction
        analysis = await self._call_openai_analysis(raw_text)

        # 3. Save Resume record
        extracted_data = {
            "skills": analysis.get("skills", []),
            "education": analysis.get("education", []),
            "experience": analysis.get("experience", []),
            "certifications": analysis.get("certifications", []),
            "projects": analysis.get("projects", []),
        }

        # Simulated path
        file_path = f"uploads/{user_id}/{filename}"

        resume_obj = Resume(
            user_id=user_id,
            filename=filename,
            file_path=file_path,
            raw_text=raw_text,
            extracted_data=extracted_data
        )
        self.db.add(resume_obj)
        await self.db.flush()  # Populates resume_obj.id

        # 4. Save Resume Scores
        score_obj = ResumeScore(
            resume_id=resume_obj.id,
            overall_score=analysis.get("overall_score", 70),
            structure_score=analysis.get("structure_score", 70),
            content_score=analysis.get("content_score", 70),
            suggestions=analysis.get("suggestions", {"structure": [], "content": []})
        )
        self.db.add(score_obj)

        # 5. Save ATS Report
        ats_obj = ATSReport(
            resume_id=resume_obj.id,
            ats_score=analysis.get("ats_score", 70),
            missing_keywords=analysis.get("missing_keywords", []),
            formatting_issues=analysis.get("formatting_issues", []),
            relevance_score=analysis.get("relevance_score", 70)
        )
        self.db.add(ats_obj)

        # Commit transaction
        await self.db.commit()
        await self.db.refresh(resume_obj)
        await self.db.refresh(score_obj)
        await self.db.refresh(ats_obj)

        return resume_obj, score_obj, ats_obj
